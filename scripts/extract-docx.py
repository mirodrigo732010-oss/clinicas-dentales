#!/usr/bin/env python3
"""
Extract text from Word (.docx) file.
Usage: python3 extract-docx.py <docx_file_path>
"""

import sys
import os

def extract_docx(docx_path: str) -> str:
    """Extract text from a Word document."""
    try:
        from docx import Document
    except ImportError:
        print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    if not os.path.exists(docx_path):
        print(f"Error: File not found: {docx_path}", file=sys.stderr)
        sys.exit(1)

    try:
        doc = Document(docx_path)
        
        text_parts = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        result = '\n\n'.join(text_parts)
        
        # Clean up
        result = result.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive whitespace
        lines = result.split('\n')
        cleaned_lines = []
        prev_empty = False
        
        for line in lines:
            stripped = line.strip()
            if not stripped:
                if not prev_empty:
                    cleaned_lines.append('')
                prev_empty = True
            else:
                cleaned_lines.append(stripped)
                prev_empty = False
        
        return '\n'.join(cleaned_lines)
        
    except Exception as e:
        print(f"Error processing DOCX: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 extract-docx.py <docx_file_path>", file=sys.stderr)
        sys.exit(1)
    
    docx_path = sys.argv[1]
    text = extract_docx(docx_path)
    print(text)
